#!/usr/bin/env python3
"""
文档转换脚本 - 支持 Word 和 PDF 转 Markdown
用法:
    python convert.py input.docx output.md
    python convert.py input.pdf output.md
    python convert.py --batch /input/dir /output/dir
"""

import argparse
import os
import sys
import subprocess
from pathlib import Path


def check_dependencies():
    """检查所需依赖是否已安装"""
    deps = {
        'pandoc': 'brew install pandoc (用于 Word 转换)',
        'pdftotext': 'brew install poppler (用于 PDF 转换)',
        'tesseract': 'brew install tesseract (用于 OCR)',
    }

    available = {}
    for cmd, install_hint in deps.items():
        try:
            subprocess.run(['which', cmd], check=True, capture_output=True)
            available[cmd] = True
        except subprocess.CalledProcessError:
            available[cmd] = False
            print(f"警告：{cmd} 未安装 - {install_hint}")

    return available


def convert_docx_to_md(input_path: str, output_path: str) -> bool:
    """
    将 Word 文档转换为 Markdown

    优先使用 pandoc，如果不可用则使用 Python 库
    """
    input_file = Path(input_path)
    output_file = Path(output_path)

    if not input_file.exists():
        print(f"错误：文件不存在 {input_path}")
        return False

    # 方法 1: 使用 pandoc (推荐)
    try:
        result = subprocess.run([
            'pandoc',
            str(input_file),
            '-f', 'docx',
            '-t', 'markdown',
            '-o', str(output_file)
        ], capture_output=True, text=True, timeout=60)

        if result.returncode == 0:
            print(f"成功使用 pandoc 转换：{output_path}")
            return True
        else:
            print(f"pandoc 转换失败：{result.stderr}")
    except FileNotFoundError:
        print("pandoc 未安装，尝试使用 Python 库...")
    except subprocess.TimeoutExpired:
        print("pandoc 转换超时")

    # 方法 2: 使用 python-docx
    try:
        from docx import Document
        import re

        doc = Document(str(input_file))
        md_content = []

        for para in doc.paragraphs:
            text = para.text
            style = para.style.name if para.style else ''

            # 处理标题
            if style.startswith('Heading'):
                level = int(style.replace('Heading ', '')) if style != 'Heading' else 1
                md_content.append(f"{'#' * level} {text}\n")
            # 处理列表
            elif style.startswith('List'):
                md_content.append(f"- {text}\n")
            # 普通段落
            elif text:
                md_content.append(f"{text}\n")

        # 处理表格
        for table in doc.tables:
            md_content.append("\n")
            for row in table.rows:
                cells = [cell.text.replace('|', '\\|') for cell in row.cells]
                md_content.append('| ' + ' | '.join(cells) + ' |\n')
                md_content.append('|' + '|'.join(['---'] * len(cells)) + '|\n')
            md_content.append("\n")

        output_file.write_text('\n'.join(md_content))
        print(f"成功使用 python-docx 转换：{output_path}")
        return True

    except ImportError:
        print("错误：需要安装 python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"转换失败：{e}")
        return False


def convert_pdf_to_md(input_path: str, output_path: str, use_ocr: bool = False) -> bool:
    """
    将 PDF 转换为 Markdown

    文字型 PDF 使用 pdftotext，扫描版使用 OCR
    """
    input_file = Path(input_path)
    output_file = Path(output_path)

    if not input_file.exists():
        print(f"错误：文件不存在 {input_path}")
        return False

    # 方法 1: 使用 pdftotext (推荐，速度快)
    try:
        result = subprocess.run([
            'pdftotext',
            '-layout',  # 保持布局
            str(input_file),
            str(output_file)
        ], capture_output=True, text=True, timeout=60)

        if result.returncode == 0:
            # 后处理：将文本转换为更规范的 Markdown
            content = output_file.read_text()
            md_content = postprocess_pdf_text(content)
            output_file.write_text(md_content)
            print(f"成功使用 pdftotext 转换：{output_path}")
            return True
        else:
            print(f"pdftotext 转换失败：{result.stderr}")
    except FileNotFoundError:
        print("pdftotext 未安装，尝试使用 pymupdf...")
    except subprocess.TimeoutExpired:
        print("pdftotext 转换超时")

    # 方法 2: 使用 pymupdf
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(input_file))
        md_content = []

        for page_num, page in enumerate(doc):
            text = page.get_text()
            md_content.append(f"<!-- Page {page_num + 1} -->\n")
            md_content.append(postprocess_pdf_text(text))
            md_content.append("\n\n---\n\n")

        doc.close()
        output_file.write_text('\n'.join(md_content))
        print(f"成功使用 pymupdf 转换：{output_path}")
        return True

    except ImportError:
        print("pymupdf 未安装，尝试 OCR 模式...")
    except Exception as e:
        print(f"pymupdf 转换失败：{e}")

    # 方法 3: OCR (用于扫描版 PDF)
    if use_ocr:
        return convert_pdf_ocr(input_path, output_path)

    return False


def convert_pdf_ocr(input_path: str, output_path: str) -> bool:
    """
    使用 OCR 识别扫描版 PDF
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract

        # 将 PDF 转为图片
        images = convert_from_path(input_path)

        md_content = []
        for i, image in enumerate(images):
            md_content.append(f"<!-- Page {i + 1} -->\n")
            # OCR 识别
            text = pytesseract.image_to_string(image)
            md_content.append(text)
            md_content.append("\n\n---\n\n")

        output_file = Path(output_path)
        output_file.write_text('\n'.join(md_content))
        print(f"成功使用 OCR 转换：{output_path}")
        return True

    except ImportError as e:
        print(f"错误：OCR 需要安装依赖：pip install pdf2image pytesseract")
        print(f"还需要安装 tesseract: brew install tesseract")
        return False
    except Exception as e:
        print(f"OCR 转换失败：{e}")
        return False


def postprocess_pdf_text(text: str) -> str:
    """
    后处理 PDF 提取的文本，转换为更规范的 Markdown
    """
    lines = text.split('\n')
    result = []
    prev_blank = True

    for line in lines:
        line = line.rstrip()

        # 检测可能的标题 (全大写或短行)
        if line and len(line) < 100 and line.isupper() and not line.endswith('.'):
            result.append(f"\n## {line.title()}\n")
            prev_blank = True
            continue

        # 检测空行
        if not line:
            if not prev_blank:
                result.append('')
                prev_blank = True
            continue

        # 普通文本
        result.append(line)
        prev_blank = False

    return '\n'.join(result)


def batch_convert(input_dir: str, output_dir: str, target_ext: str = '.md') -> dict:
    """
    批量转换目录中的所有文档

    Returns:
        dict: {'success': [...], 'failed': [...]}
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    results = {'success': [], 'failed': []}

    # 查找所有支持的文档
    supported_exts = ['.docx', '.pdf']
    files = [f for f in input_path.iterdir() if f.suffix.lower() in supported_exts]

    print(f"找到 {len(files)} 个文件待转换...\n")

    for i, file in enumerate(files):
        output_name = file.stem + target_ext
        output_file = output_path / output_name

        print(f"[{i+1}/{len(files)}] 转换：{file.name}")

        if file.suffix.lower() == '.docx':
            success = convert_docx_to_md(str(file), str(output_file))
        elif file.suffix.lower() == '.pdf':
            success = convert_pdf_to_md(str(file), str(output_file))
        else:
            success = False

        if success:
            results['success'].append(file.name)
        else:
            results['failed'].append(file.name)

        print()

    # 打印摘要
    print("=" * 50)
    print(f"转换完成：成功 {len(results['success'])}, 失败 {len(results['failed'])}")
    if results['failed']:
        print(f"失败文件：{', '.join(results['failed'])}")

    return results


def main():
    parser = argparse.ArgumentParser(
        description='文档转换工具 - 支持 Word 和 PDF 转 Markdown'
    )
    parser.add_argument('input', help='输入文件或目录路径')
    parser.add_argument('output', nargs='?', help='输出文件或目录路径')
    parser.add_argument('--ocr', action='store_true', help='对 PDF 使用 OCR 识别')
    parser.add_argument('--batch', action='store_true', help='批量转换目录')

    args = parser.parse_args()

    # 检查依赖
    print("检查依赖...")
    deps = check_dependencies()
    print()

    # 批量转换模式
    if args.batch:
        input_dir = args.input
        output_dir = args.output or f"{input_dir}_md"
        batch_convert(input_dir, output_dir)
        return

    # 单文件转换
    if not args.output:
        print("错误：需要指定输出文件路径")
        print("用法：python convert.py input.docx output.md")
        print("      python convert.py --batch /input/dir /output/dir")
        sys.exit(1)

    input_file = Path(args.input)
    if not input_file.exists():
        print(f"错误：文件不存在 {args.input}")
        sys.exit(1)

    # 根据扩展名选择转换方式
    if input_file.suffix.lower() == '.docx':
        success = convert_docx_to_md(args.input, args.output)
    elif input_file.suffix.lower() == '.pdf':
        success = convert_pdf_to_md(args.input, args.output, use_ocr=args.ocr)
    else:
        print(f"不支持的文件格式：{input_file.suffix}")
        print("支持的格式：.docx, .pdf")
        success = False

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
